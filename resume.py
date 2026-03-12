# Create a professionally formatted DOCX resume for the user using python-docx
from docx import Document

doc = Document()

doc.add_heading('MANISH MISHRA', 0)
p = doc.add_paragraph()
p.add_run("Email: ").bold = True
p.add_run("manishkr.mishra27d@gmail.com\n")
p.add_run("Phone: ").bold = True
p.add_run("(+91) 9643470624\n")
p.add_run("LinkedIn: ").bold = True
p.add_run("https://www.linkedin.com/in/manish-mishra-647518172")

doc.add_heading('PROFESSIONAL SUMMARY', level=1)
doc.add_paragraph(
"Application Support Engineer with 4.7 years of experience in production support and application "
"operations within BFSI domain at Tata Consultancy Services. Skilled in incident management, "
"API troubleshooting, Unix/Linux systems, SQL database support, and job scheduling tools. "
"Experienced in ITIL-based support models, monitoring tools, and automation of operational tasks. "
"Adept at supporting critical banking applications, performing root cause analysis, and ensuring "
"high system availability in production environments."
)

doc.add_heading('TECHNICAL SKILLS', level=1)

skills = {
"Application & Production Support": "Incident Management, Problem Management, Change Management (ITIL), Production Support, Run the Bank operations",
"Programming & Scripting": "Shell Scripting, Basic Python scripting, Perl (basic understanding)",
"Operating Systems": "Unix / Linux, Basic Windows Server Administration",
"Databases": "SQL / PL-SQL, MS SQL Server, Oracle, MariaDB, Amazon RDS",
"Job Scheduling Tools": "Autosys, IBM Tivoli Workload Scheduler (TWS)",
"Monitoring Tools": "AppDynamics, ITRS Geneos, Grafana, Wiley Introscope",
"DevOps Exposure": "CI/CD support, Nolio, Chef",
"Big Data": "Hadoop (HDFS, MapReduce), Hadoop cluster monitoring, HBase management",
"ETL": "AbInitio / Informatica job monitoring and troubleshooting",
"Cloud": "Understanding of On-Prem, Public Cloud, Private Cloud, SaaS environments"
}

for k,v in skills.items():
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(f"{k}: ").bold = True
    para.add_run(v)

doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)

doc.add_paragraph("L2 APPLICATION SUPPORT ENGINEER – BFSI OPERATIONS", style='List Bullet')
doc.add_paragraph("Tata Consultancy Services | 2021 – Present")

resp = [
"Provide L2 production support for BFSI applications ensuring high availability.",
"Manage incidents, problems, and change requests following ITIL practices.",
"Perform Unix/Linux troubleshooting, log analysis, and service management.",
"Support React-based applications by analyzing API failures and backend responses.",
"Monitor applications using AppDynamics, Geneos, and Grafana.",
"Manage and troubleshoot batch jobs using Autosys/TWS.",
"Execute SQL/PLSQL queries for data validation and issue investigation.",
"Analyze database performance issues and execution plans.",
"Support ETL workflows using AbInitio/Informatica.",
"Assist in automation of manual support activities using shell and Python scripts.",
"Work with DevOps teams on CI/CD pipelines and automation tools such as Nolio and Chef.",
"Support Hadoop ecosystem jobs and data verification in HDFS environments."
]

for r in resp:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('KEY ACHIEVEMENTS', level=1)

achievements = [
"Automated repetitive operational tasks using shell scripting.",
"Handled critical production incidents and minimized application downtime.",
"Improved monitoring alerts and dashboards for faster issue detection.",
"Stabilized batch job failures and ETL workflows in production environments."
]

for a in achievements:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('EDUCATION', level=1)
doc.add_paragraph("Bachelor’s Degree in Engineering / Technology (Add University & Year)")

doc.add_heading('ADDITIONAL INFORMATION', level=1)
doc.add_paragraph("Strong communication and stakeholder management skills.")
doc.add_paragraph("Experience working in 24x7 production support environments.")
doc.add_paragraph("Flexible to work in shift rotations and weekend on-call support.")

file_path = "Manish_Mishra_Application_Support_Resume.docx"
doc.save(file_path)

file_path